"""
RAG Chunker for Post-Op Patient Q&A
─────────────────────────────────────────────
Recursive chunking: section-aware paragraph splitting that keeps numbered
lists, instructions and related text together for better completeness.

Optimised for throughput:
  • OpenRouter embeddings are sent as true batch requests
  • Multiple documents are processed concurrently
  • Supports: PDF, TXT, MD, DOCX
"""
import asyncio
import hashlib
import json
import os
import re
import time
import requests
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any
import logging
from tqdm.asyncio import tqdm

import fitz  # PyMuPDF
from langchain_core.documents import Document

import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '../../..')))

from app.core.config import settings
from app.core.factory import ModelFactory
import pickle

# ── Logging ──────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)

# ระงับ Log ที่ไม่จำเป็นจากไลบรารีเบื้องหลังระหว่างเรียก API
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("httpcore").setLevel(logging.WARNING)
logging.getLogger("chromadb").setLevel(logging.ERROR)
logging.getLogger("langchain").setLevel(logging.WARNING)

# ── Tunables ─────────────────────────────────────────────────────────
MAX_CONCURRENT_EMBED = 5     # parallel embedding batch requests
EMBED_BATCH_SIZE     = 30    # ลดเหลือ 30 texts per call เพื่อป้องกัน Timeout จากฝั่ง OpenRouter
CHROMA_ADD_BATCH     = 200   # docs per Chroma add_documents call

# ── Recursive Chunking Tunables ──────────────────────────────────────
CHUNK_MAX_CHARS = 1000  # max characters per chunk (keeps numbered lists intact)
CHUNK_OVERLAP   = 150   # overlap characters between consecutive chunks

# ── Paths ────────────────────────────────────────────────────────────
DOCUMENT_DIR    = settings.DATA_DIR / "document"
CHROMADB_DIR    = settings.DATA_DIR / "chroma_db"   # used for BM25/manifest only
COLLECTION_NAME = "post_op_propositions"


def _get_chroma_client():
    """Return Chroma client.

    Priority:
      1. Chroma Cloud — if CHROMA_API_KEY + CHROMA_TENANT + CHROMA_DATABASE are set
      2. HttpClient   — if CHROMA_HOST is reachable (Docker / self-hosted)
      3. PersistentClient — local fallback (dev / single-server deploy)
    """
    import chromadb as _c

    api_key  = getattr(settings, "CHROMA_API_KEY", "")
    tenant   = getattr(settings, "CHROMA_TENANT", "")
    database = getattr(settings, "CHROMA_DATABASE", "")

    if api_key and tenant and database:
        client = _c.CloudClient(tenant=tenant, database=database, api_key=api_key)
        logger.info("[chroma] Using Chroma Cloud")
        return client

    host = getattr(settings, "CHROMA_HOST", "localhost")
    port = int(getattr(settings, "CHROMA_PORT", 8002))
    try:
        client = _c.HttpClient(host=host, port=port)
        client.heartbeat()
        logger.info(f"[chroma] Using HttpClient at {host}:{port}")
        return client
    except Exception:
        logger.warning("[chroma] HttpClient unreachable — falling back to PersistentClient")
        CHROMADB_DIR.mkdir(parents=True, exist_ok=True)
        return _c.PersistentClient(path=str(CHROMADB_DIR))

MANIFEST_FILE = CHROMADB_DIR / "doc_manifest.json"
STATUS_FILE   = CHROMADB_DIR / "ingest_status.json"

# ── BM25 Thai Tokenizer ─────────────────────────────────────────────
from app.services.common.thai_tokenizer import bm25_thai_preprocess  # noqa: F401

# ── Supported file types ─────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


# ── Embeddings  ─────────────────────────────────────────────────────────

# Note: The OpenRouterEmbeddings wrapper is kept here for reference or 
# if you need to use it directly, but the factory is the preferred way 
# to get embeddings.

class OpenRouterEmbeddings:
    """Wrapper for OpenRouter embedding models with batch support."""

    def __init__(self, model: str = "BAAI/bge-m3"):
        self.model = model
        self.api_key = os.getenv("OPEN_ROUTER_API_KEY", "")
        if not self.api_key:
            logger.warning("OPEN_ROUTER_API_KEY is not set. Embeddings might fail.")

    def _get_embeddings_batch(self, texts: List[str]) -> List[List[float]]:
        import time
        max_retries = 3
        base_delay = 2

        for attempt in range(max_retries):
            try:
                response = requests.post(
                    url="https://openrouter.ai/api/v1/embeddings",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": self.model,
                        "input": texts,
                        "encoding_format": "float",
                    },
                    timeout=90,  # เพิ่ม Timeout เป็น 90 วินาที
                )
                response.raise_for_status()
                data = response.json()["data"]
                data.sort(key=lambda x: x["index"])
                return [d["embedding"] for d in data]
            
            except Exception as e:
                logger.error(f"Batch embedding error (Attempt {attempt+1}/{max_retries}): {e}")
                if attempt < max_retries - 1:
                    sleep_time = base_delay * (2 ** attempt)
                    logger.info(f"Retrying embedding batch in {sleep_time} seconds...")
                    time.sleep(sleep_time)
                else:
                    logger.error("Max retries reached. Returning zero embeddings.")
                    return [[0.0] * 1024 for _ in texts]

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from tqdm import tqdm

        total_batches = (len(texts) + EMBED_BATCH_SIZE - 1) // EMBED_BATCH_SIZE

        # Build ordered list of (batch_index, batch_texts)
        batches = [
            (i // EMBED_BATCH_SIZE, texts[i : i + EMBED_BATCH_SIZE])
            for i in range(0, len(texts), EMBED_BATCH_SIZE)
        ]

        results: dict[int, List[List[float]]] = {}
        with ThreadPoolExecutor(max_workers=MAX_CONCURRENT_EMBED) as executor:
            futures = {
                executor.submit(self._get_embeddings_batch, batch): idx
                for idx, batch in batches
            }
            with tqdm(total=total_batches, desc="Embedding Batches",
                      unit="batch", dynamic_ncols=True) as pbar:
                for future in as_completed(futures):
                    idx = futures[future]
                    results[idx] = future.result()
                    pbar.update(1)
                    pbar.set_postfix(texts=f"{min((idx+1)*EMBED_BATCH_SIZE, len(texts))}/{len(texts)}")

        # Reassemble in original order
        all_embeddings: List[List[float]] = []
        for i in range(total_batches):
            all_embeddings.extend(results[i])
        return all_embeddings

    def embed_query(self, text: str) -> List[float]:
        return self._get_embeddings_batch([text])[0]


# =====================================================================
#  Text Extraction (multi-format)
# =====================================================================
def extract_text_from_pdf(pdf_path: str) -> List[Dict[str, Any]]:
    """Extract text from PDF with basic heading detection."""
    doc = fitz.open(pdf_path)
    content = []
    current_heading = "General"

    for page in doc:
        blocks = page.get_text("dict")["blocks"]
        for b in blocks:
            if "lines" in b:
                for line in b["lines"]:
                    for s in line["spans"]:
                        text = s["text"].strip()
                        if not text:
                            continue
                        is_heading = s["size"] > 12 or (text.isupper() and len(text) < 50)
                        if is_heading:
                            current_heading = text
                        else:
                            content.append({
                                "text": text,
                                "heading": current_heading,
                                "page": page.number + 1,
                            })
    return content


def extract_text_from_txt(txt_path: str) -> List[Dict[str, Any]]:
    """Extract text from plain text file with basic heading detection."""
    with open(txt_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    content = []
    current_heading = "General"

    for line in lines:
        text = line.strip()
        if not text:
            continue
        # Heuristic: short lines that are UPPERCASE or end with ":" are headings
        is_heading = (
            (text.isupper() and len(text) < 80)
            or (len(text) < 80 and text.endswith(":"))
        )
        if is_heading:
            current_heading = text.rstrip(":")
        else:
            content.append({"text": text, "heading": current_heading})

    return content


def extract_text_from_markdown(md_path: str) -> List[Dict[str, Any]]:
    """Extract text from Markdown file using '#' headings.
    
    Includes cleaning for scraped web pages: strips navigation headers,
    footers, contact info, form fields, and other boilerplate.
    """
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # ── Phase 1: Find content boundaries ──────────────────────────────
    # Skip navigation header: everything before the first H1 (#) heading
    first_h1_idx = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# ") and not stripped.startswith("## "):
            first_h1_idx = i
            break

    start_idx = first_h1_idx if first_h1_idx is not None else 0

    # Find footer boundary: last '---' separator or common footer markers
    end_idx = len(lines)
    footer_markers = {"Source:", "Scraped:", "---"}
    # Common footer section headings (exact match after stripping #)
    footer_headings = {
        "แชร์", "หากสนใจต้องการปรึกษาแพทย์", "สอบถามข้อมูลเพิ่มเติม",
        "บทความที่คุณอาจสนใจ", "บทความที่เกี่ยวข้อง",
    }

    # Scan from the content area to find where footer begins
    content_started = False
    for i in range(start_idx, len(lines)):
        stripped = lines[i].strip()
        if not stripped:
            continue

        # Mark that we've seen real content (past the heading and metadata)
        if not content_started and not stripped.startswith("#") and len(stripped) > 30:
            content_started = True

        if content_started:
            # Detect footer start: '---' after content, or source/scraped lines
            if stripped == "---":
                # Check if there's still real content after this '---'
                has_content_after = False
                for j in range(i + 1, min(i + 10, len(lines))):
                    future = lines[j].strip()
                    if future and not future.startswith("#") and len(future) > 50:
                        if not any(future.startswith(m) for m in footer_markers):
                            has_content_after = True
                            break
                if not has_content_after:
                    end_idx = i
                    break

            if stripped.startswith("Source:") or stripped.startswith("Scraped:"):
                end_idx = i
                break

            # Detect footer heading sections
            if stripped.startswith("#"):
                heading_text = stripped.lstrip("# ").strip()
                if heading_text in footer_headings:
                    end_idx = i
                    break

    # ── Phase 2: Extract content with heading tracking ────────────────
    # Patterns to skip (junk lines within content area)
    _JUNK_PATTERNS = {
        "ค้นหา", "เข้าสู่ระบบ", "แชร์", "ส่งข้อความ", "ยืนยัน",
        "ชื่อ", "นามสกุล", "เบอร์ติดต่อ", "อีเมล", "รายละเอียด",
        "ที่อยู่", "เวลาทำการ", "เบอร์โทร",
    }
    # Regex patterns for common footer/boilerplate lines
    _JUNK_RE = re.compile(
        r'^(สอบถาม(รายละเอียด|ข้อมูล)|รับข่าวสาร|โทร\s*\d|Facebook\s*:|Line\s*(official|:)|'
        r'ปรึกษาปัญหาเรื่องฟัน|ดูโปรโมชั่น|กรุณากรอกข้อมูล)',
        re.IGNORECASE
    )

    content = []
    current_heading = "General"

    for line in lines[start_idx:end_idx]:
        text = line.strip()
        if not text:
            continue

        # Detect Markdown headings
        if text.startswith("#"):
            heading_text = text.lstrip("# ").strip()
            # Remove bold markers from heading text
            heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
            # Skip navigation-level headings (###### typically = nav)
            if text.startswith("######"):
                continue
            # Skip junk headings (social share, forms, etc.)
            if heading_text in _JUNK_PATTERNS or heading_text in footer_headings:
                continue
            current_heading = heading_text
            continue

        # Skip standalone markdown links like [TH](/)[EN](/en) or [ดูทั้งหมด](...)
        if re.match(r'^\[.+?\]\(.+?\)(\[.+?\]\(.+?\))*$', text):
            continue

        # Skip bullet items that are just metadata (date, branch name)
        if text.startswith("* ") and len(text) < 30:
            continue

        # Clean markdown formatting FIRST (before junk detection)
        # Remove bold markers **text** → text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # Remove inline links [text](url) → text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        # Strip trailing whitespace (markdown line breaks)
        text = text.rstrip()

        if not text or len(text) <= 3:
            continue

        # Skip junk lines (after markdown cleaning)
        if text in _JUNK_PATTERNS:
            continue
        if _JUNK_RE.search(text):
            continue

        content.append({"text": text, "heading": current_heading})

    return content


def extract_text_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """Extract text from DOCX file using paragraph styles."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return []

    doc = DocxDocument(docx_path)
    content = []
    current_heading = "General"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Detect heading styles (Heading 1, Heading 2, etc.)
        style_name = (para.style.name or "").lower()
        if "heading" in style_name or "title" in style_name:
            current_heading = text
        else:
            content.append({"text": text, "heading": current_heading})

    return content


def extract_text_from_file(file_path: Path) -> List[Dict[str, Any]]:
    """Route to correct extractor based on file extension."""
    ext = file_path.suffix.lower()
    extractors = {
        ".pdf":  lambda p: extract_text_from_pdf(str(p)),
        ".txt":  lambda p: extract_text_from_txt(str(p)),
        ".md":   lambda p: extract_text_from_markdown(str(p)),
        ".docx": lambda p: extract_text_from_docx(str(p)),
    }
    extractor = extractors.get(ext)
    if extractor is None:
        logger.warning(f"Unsupported file type: {ext} — skipping {file_path.name}")
        return []
    return extractor(file_path)


# =====================================================================
#  Recursive chunking  (section-aware, preserves lists/paragraphs)
# =====================================================================
def _split_section_with_overlap(
    lines: List[str],
    max_chars: int = CHUNK_MAX_CHARS,
    overlap_chars: int = CHUNK_OVERLAP,
) -> List[str]:
    """Split a list of lines into chunks at line boundaries with overlap."""
    chunks: List[str] = []
    current_lines: List[str] = []
    current_len = 0

    for line in lines:
        line_len = len(line) + 1  # +1 for newline
        if current_len + line_len > max_chars and current_lines:
            chunks.append("\n".join(current_lines))
            # Build overlap from end of current chunk
            overlap_lines: List[str] = []
            overlap_len = 0
            for prev_line in reversed(current_lines):
                if overlap_len + len(prev_line) + 1 > overlap_chars:
                    break
                overlap_lines.insert(0, prev_line)
                overlap_len += len(prev_line) + 1
            current_lines = overlap_lines
            current_len = overlap_len
        current_lines.append(line)
        current_len += line_len

    if current_lines:
        chunks.append("\n".join(current_lines))

    return chunks


def recursive_chunk_document(file_path: Path) -> List[Document]:
    """Chunk a document into sections recursively.

    Groups text by heading, keeps complete sections where possible,
    and splits large sections at line boundaries with overlap.
    """
    filename = file_path.name
    extracted = extract_text_from_file(file_path)

    if not extracted:
        return []

    # Group items by heading, preserving order
    sections: list[tuple[str, list[str]]] = []
    current_heading: str | None = None
    current_texts: list[str] = []

    for item in extracted:
        if item["heading"] != current_heading:
            if current_heading is not None and current_texts:
                sections.append((current_heading, current_texts))
            current_heading = item["heading"]
            current_texts = []
        current_texts.append(item["text"])

    if current_heading is not None and current_texts:
        sections.append((current_heading, current_texts))

    documents: List[Document] = []
    for heading, texts in sections:
        section_text = "\n".join(texts)

        if len(section_text) <= CHUNK_MAX_CHARS:
            header = f"[SOURCE: {filename} | SECTION: {heading}]\n"
            documents.append(Document(
                page_content=header + section_text,
                metadata={"source": filename, "section": heading},
            ))
        else:
            for chunk_text in _split_section_with_overlap(texts):
                header = f"[SOURCE: {filename} | SECTION: {heading}]\n"
                documents.append(Document(
                    page_content=header + chunk_text,
                    metadata={"source": filename, "section": heading},
                ))

    logger.debug(f"  Recursive: {filename} → {len(documents)} chunks")
    return documents


# =====================================================================
#  Document-level processing  (all batches in parallel)
# =====================================================================
async def process_document(file_path: Path) -> List[Document]:
    """Process a single document (PDF/TXT/MD/DOCX) using recursive chunking."""
    filename = file_path.name
    logger.debug(f"Processing: {filename}")
    return recursive_chunk_document(file_path)


# =====================================================================
#  Manifest + Status helpers
# =====================================================================
def compute_file_hash(file_path: Path) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def load_manifest() -> dict:
    if MANIFEST_FILE.exists():
        with open(MANIFEST_FILE) as f:
            return json.load(f)
    return {}


def save_manifest(manifest: dict) -> None:
    with open(MANIFEST_FILE, "w") as f:
        json.dump(manifest, f, indent=2, ensure_ascii=False)


def write_status(status: str, mode: str = "", message: str = "",
                 files_processed: int = 0, files_total: int = 0, error: str = "") -> None:
    CHROMADB_DIR.mkdir(parents=True, exist_ok=True)
    with open(STATUS_FILE, "w") as f:
        json.dump({
            "status": status,
            "mode": mode,
            "message": message,
            "files_processed": files_processed,
            "files_total": files_total,
            "updated_at": datetime.utcnow().isoformat(),
            "error": error,
        }, f)


def read_status() -> dict:
    if STATUS_FILE.exists():
        with open(STATUS_FILE) as f:
            return json.load(f)
    return {"status": "idle"}


# =====================================================================
#  Main ingestion pipeline
# =====================================================================
async def ingest(mode: str = "reingest", *, reset: bool = False):
    """Ingest documents into ChromaDB + BM25.

    mode: "reingest" — delete collection, process all files, rebuild from scratch
          "append"   — hash-based incremental: only new/changed files
    reset=True is kept for CLI backwards compat and maps to mode="reingest".
    """
    if reset:
        mode = "reingest"

    t0 = time.perf_counter()
    write_status("running", mode=mode, message="Starting…")
    logger.info(f"═══ Starting ingestion ({mode}) ═══")

    try:
        import chromadb as _chromadb
        from chromadb.errors import NotFoundError

        # ── 1. Sync files from Supabase Storage → local DOCUMENT_DIR ──
        DOCUMENT_DIR.mkdir(parents=True, exist_ok=True)
        try:
            from app.services.storage_service import list_files as storage_list, download_to_temp
            storage_files = storage_list()
            for sf in storage_files:
                name = sf.get("name", "")
                if Path(name).suffix.lower() in SUPPORTED_EXTENSIONS:
                    local_path = DOCUMENT_DIR / name
                    if not local_path.exists():
                        logger.info(f"  ↓ Downloading {name} from Supabase Storage…")
                        download_to_temp(name, DOCUMENT_DIR)
        except Exception as e:
            logger.warning(f"  ⚠ Supabase Storage sync skipped: {e}")

        # ── 2. Determine files to process ────────────────────────────
        all_files: list[Path] = []
        for ext in SUPPORTED_EXTENSIONS:
            all_files.extend(DOCUMENT_DIR.rglob(f"*{ext}"))

        manifest = load_manifest()

        if mode == "reingest":
            # Delete only the RAG collection
            client = _get_chroma_client()
            try:
                client.delete_collection(name=COLLECTION_NAME)
                logger.info(f"  ✓ Collection '{COLLECTION_NAME}' deleted.")
            except (ValueError, NotFoundError):
                pass
            manifest = {}
            files_to_process = all_files
        else:
            # append: compute hashes, find new/changed/deleted
            current_hashes: dict[str, str] = {}
            for f in all_files:
                try:
                    current_hashes[f.name] = compute_file_hash(f)
                except Exception as e:
                    logger.warning(f"Cannot hash {f.name}: {e}")

            deleted_names = {n for n in manifest if n not in current_hashes}
            changed_names = {n for n, h in current_hashes.items()
                             if n in manifest and manifest[n].get("sha256") != h}
            new_names     = {n for n in current_hashes if n not in manifest}

            # Remove stale chunks from ChromaDB
            stale = deleted_names | changed_names
            if stale:
                client = _get_chroma_client()
                try:
                    col = client.get_collection(COLLECTION_NAME)
                    for fname in stale:
                        try:
                            col.delete(where={"source": fname})
                            logger.info(f"  Removed stale chunks for: {fname}")
                        except Exception:
                            pass
                except (ValueError, NotFoundError):
                    pass
                for fname in stale:
                    manifest.pop(fname, None)

            files_to_process = [f for f in all_files if f.name in (new_names | changed_names)]
            if not files_to_process:
                write_status("done", mode=mode, message="Nothing changed — index is up to date.",
                             files_processed=0, files_total=0)
                logger.info("No changed files — index is up to date.")
                save_manifest(manifest)
                return

        write_status("running", mode=mode, message=f"Chunking {len(files_to_process)} file(s)…",
                     files_processed=0, files_total=len(files_to_process))

        if not files_to_process:
            write_status("done", mode=mode, message="No files to process.", files_processed=0, files_total=0)
            return

        logger.info(f"Processing {len(files_to_process)} file(s)…")

        # ── 2. Chunk ──────────────────────────────────────────────────
        async def safe_process(f: Path):
            try:
                return await process_document(f)
            except Exception as e:
                logger.error(f"Error processing {f.name}: {e}")
                return []

        doc_results = await tqdm.gather(*[safe_process(f) for f in files_to_process],
                                        desc="Chunking")
        new_docs: list[Document] = []
        for r in doc_results:
            if isinstance(r, list):
                new_docs.extend(r)

        if not new_docs:
            write_status("done", mode=mode, message="No chunks generated.", files_processed=len(files_to_process), files_total=len(files_to_process))
            return

        write_status("running", mode=mode, message=f"Embedding {len(new_docs)} chunks…",
                     files_processed=0, files_total=len(files_to_process))

        # ── 3. Embed ──────────────────────────────────────────────────
        embedding_function = ModelFactory.get_embeddings(use_case="rag")
        texts = [doc.page_content for doc in new_docs]
        all_embeddings = embedding_function.embed_documents(texts)

        # ── 4. Store in ChromaDB ──────────────────────────────────────
        import uuid as _uuid
        from langchain_chroma import Chroma

        vectorstore = Chroma(
            collection_name=COLLECTION_NAME,
            client=_get_chroma_client(),
            embedding_function=embedding_function,
        )
        logger.info(f"Adding {len(new_docs)} chunks to ChromaDB…")
        for i in range(0, len(new_docs), CHROMA_ADD_BATCH):
            batch_docs = new_docs[i: i + CHROMA_ADD_BATCH]
            batch_embs = all_embeddings[i: i + CHROMA_ADD_BATCH]
            vectorstore._collection.add(
                ids=[str(_uuid.uuid4()) for _ in batch_docs],
                embeddings=batch_embs,
                documents=[d.page_content for d in batch_docs],
                metadatas=[d.metadata for d in batch_docs],
            )
        logger.info("✓ ChromaDB updated.")

        # ── 5. Rebuild BM25 from ALL docs in collection ───────────────
        write_status("running", mode=mode, message="Rebuilding BM25 index…",
                     files_processed=len(files_to_process), files_total=len(files_to_process))
        from langchain_community.retrievers import BM25Retriever

        all_result = vectorstore._collection.get(include=["documents", "metadatas"])
        all_chroma_docs = [
            Document(page_content=d, metadata=m)
            for d, m in zip(all_result["documents"], all_result["metadatas"])
        ]
        bm25_retriever = BM25Retriever.from_documents(all_chroma_docs,
                                                       preprocess_func=bm25_thai_preprocess)
        bm25_path = CHROMADB_DIR / "bm25_index.pkl"
        with open(bm25_path, "wb") as f:
            pickle.dump(bm25_retriever, f)
        logger.info(f"✓ BM25 index rebuilt with {len(all_chroma_docs)} total chunks.")

        # ── 6. Update manifest ────────────────────────────────────────
        for fp in files_to_process:
            try:
                manifest[fp.name] = {
                    "sha256": compute_file_hash(fp),
                    "size": fp.stat().st_size,
                    "last_ingested": datetime.utcnow().isoformat(),
                }
            except Exception:
                pass
        save_manifest(manifest)

        elapsed = time.perf_counter() - t0
        write_status("done", mode=mode,
                     message=f"Done — {len(new_docs)} chunks from {len(files_to_process)} file(s) in {elapsed:.1f}s",
                     files_processed=len(files_to_process), files_total=len(files_to_process))
        logger.info(f"═══ Ingestion complete in {elapsed:.1f}s ═══")

    except Exception as exc:
        write_status("error", mode=mode, error=str(exc))
        logger.exception(f"Ingestion failed: {exc}")
        raise


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description="RAG Document Ingestion Pipeline")
    parser.add_argument("--mode", choices=["reingest", "append"], default="reingest",
                        help="reingest: delete and rebuild from scratch; append: only new/changed files")
    parser.add_argument("--reset", action="store_true", help="Alias for --mode=reingest")
    parser.add_argument("-y", "--yes", action="store_true", help="Skip confirmation prompt")
    args = parser.parse_args()

    mode = "reingest" if args.reset else args.mode
    print("=====================================================")
    print("RAG Ingestion")
    print("=====================================================")
    print(f"Mode: {mode}")
    if not args.yes:
        confirm = input("\nProceed? (y/n) [y]: ").strip().lower()
        if confirm not in ('', 'y', 'yes'):
            print("Cancelled.")
            sys.exit(0)
    try:
        asyncio.run(ingest(mode=mode))
    except KeyboardInterrupt:
        sys.exit(1)
